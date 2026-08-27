import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import latex2mathml.converter
from lxml import etree

XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
xslt_tree = etree.parse(XSLT_PATH)
transform = etree.XSLT(xslt_tree)

def latex_to_omml(latex_code):
    mml = latex2mathml.converter.convert(latex_code)
    tree = etree.fromstring(mml)
    omml = transform(tree)
    return omml.getroot()

def create_document():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("YEA PROJECT | Page ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(120, 120, 120)
        
        # Add Page field
        f_p_xml = f_p._p
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        f_p_xml.append(fldSimple)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(30, 30, 30)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Helper function for adding an inline math run to a paragraph
    def add_math_to_paragraph(p, latex_code):
        omml_elem = latex_to_omml(latex_code)
        tag_name = etree.QName(omml_elem.tag).localname
        if tag_name == 'oMathPara':
            for child in omml_elem:
                p._p.append(child)
        else:
            p._p.append(omml_elem)

    # Helper function for adding a numbered display equation
    def add_display_equation(latex_code, eq_number=None, space_before=5, space_after=5):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Clear borders
        tblPr = tbl._tbl.tblPr
        tblBorders = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="none"/>'
            r'  <w:left w:val="none"/>'
            r'  <w:bottom w:val="none"/>'
            r'  <w:right w:val="none"/>'
            r'  <w:insideH w:val="none"/>'
            r'  <w:insideV w:val="none"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(tblBorders)
        
        row = tbl.rows[0]
        cell_math = row.cells[0]
        cell_num = row.cells[1]
        
        cell_math.width = Inches(5.7)
        cell_num.width = Inches(0.8)
        
        cell_math.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_num.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add equation
        p_m = cell_math.paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m.paragraph_format.space_before = Pt(space_before)
        p_m.paragraph_format.space_after = Pt(space_after)
        
        omml_elem = latex_to_omml(latex_code)
        tag_name = etree.QName(omml_elem.tag).localname
        if tag_name == 'oMathPara':
            for child in omml_elem:
                p_m._p.append(child)
        else:
            p_m._p.append(omml_elem)
            
        # Add equation number
        p_n = cell_num.paragraphs[0]
        p_n.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_n.paragraph_format.space_before = Pt(space_before)
        p_n.paragraph_format.space_after = Pt(space_after)
        if eq_number:
            r = p_n.add_run(f"({eq_number})")
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.bold = True

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 44, 87)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 70, 120)

    # -------------------------------------------------------------
    # DOCUMENT HEADER & TITLE
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("A Study on the Pressure Driven Unsteady Fluid Flow Through a Vertical Channel Filled with Anisotropic Porous Material with Navier Slip Condition")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(10, 30, 70)
    
    # Subtitle / Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(14)
    r_meta = p_meta.add_run("Mathematical Formulation, Dimensionless Analysis, and Exact Analytical Solution via Laplace Transform Method\n[YEA PROJECT]")
    r_meta.font.name = 'Times New Roman'
    r_meta.font.size = Pt(10.5)
    r_meta.font.italic = True
    r_meta.font.color.rgb = RGBColor(90, 90, 90)

    # Horizontal divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_border = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="12" w:space="1" w:color="102C57"/></w:pBdr>' % nsdecls('w'))
    p_div._p.get_or_add_pPr().append(p_div_border)

    # =============================================================
    # SECTION 1: PROBLEM FORMULATION & GOVERNING EQUATIONS (PAGE 1)
    # =============================================================
    add_heading_1("1. Mathematical Formulation & Governing Equations (Page 1)")
    
    p = doc.add_paragraph("A study on the pressure-driven unsteady fluid flow through a vertical channel filled with an anisotropic porous material with Navier slip condition. The dimensional governing momentum equation is given by:")
    
    # Equation (1)
    eq1_latex = r"\frac{\partial u^\prime}{\partial t^\prime} = \nu_{\mathrm{eff}} \frac{\partial^2 u^\prime}{\partial y^{\prime 2}} - \frac{u^\prime \nu}{\bar{\bar{K}}} + G^*"
    add_display_equation(eq1_latex, eq_number="1")
    
    # Initial & Boundary Conditions (Page 1)
    p = doc.add_paragraph()
    r = p.add_run("Subject to:")
    r.font.bold = True
    
    # Initial Condition
    p_ic = doc.add_paragraph()
    p_ic.add_run("•  For ")
    add_math_to_paragraph(p_ic, r"t^\prime \le 0")
    p_ic.add_run(":")
    
    eq_ic_latex = r"u^\prime = 0 \quad \text{for } 0 \le y^\prime \le h"
    add_display_equation(eq_ic_latex)
    
    # Boundary Conditions (2)
    p_bc = doc.add_paragraph()
    p_bc.add_run("•  For ")
    add_math_to_paragraph(p_bc, r"t^\prime > 0")
    p_bc.add_run(":")
    
    eq2_latex = r"\begin{cases} u^\prime = \beta_1^* \frac{\partial u^\prime}{\partial y^\prime} & \text{at } y^\prime = 0 \\ u^\prime = -\beta_2^* \frac{\partial u^\prime}{\partial y^\prime} & \text{at } y^\prime = h \end{cases}"
    add_display_equation(eq2_latex, eq_number="2")

    # Permeability Tensor (Page 1)
    p = doc.add_paragraph()
    p.add_run("where ")
    add_math_to_paragraph(p, r"\bar{\bar{K}}")
    p.add_run(" is the symmetrical second-order permeability tensor defined as:")
    
    eq_tensor_latex = r"\bar{\bar{K}} = \begin{bmatrix} K_1 \cos^2\theta + K_2 \sin^2\theta & (K_1 - K_2)\sin\theta\cos\theta \\ (K_1 - K_2)\sin\theta\cos\theta & K_2 \cos^2\theta + K_1 \sin^2\theta \end{bmatrix}"
    add_display_equation(eq_tensor_latex)

    # Non-dimensional Quantities (Page 1)
    p = doc.add_paragraph("Introducing the non-dimensional quantities:")
    
    eq3_latex = r"u = \frac{u^\prime h}{\nu}, \quad y = \frac{y^\prime}{h}, \quad \gamma = \frac{\nu_{\mathrm{eff}}}{\nu}, \quad t = \frac{t^\prime \nu}{h^2}"
    add_display_equation(eq3_latex, eq_number="3")

    p = doc.add_paragraph("The transformations of the corresponding derivative quantities are:")
    eq_derivs_latex = r"u^\prime = \frac{\nu}{h} u, \quad \frac{\partial u^\prime}{\partial t^\prime} = \frac{\nu^2}{h^3}\frac{\partial u}{\partial t}, \quad \frac{\partial^2 u^\prime}{\partial y^{\prime 2}} = \frac{\nu}{h^3}\frac{\partial^2 u}{\partial y^2}"
    add_display_equation(eq_derivs_latex)

    # =============================================================
    # SECTION 2: NON-DIMENSIONAL MOMENTUM EQUATION & LAPLACE TRANSFORM (PAGE 2)
    # =============================================================
    add_heading_1("2. Non-Dimensional Equation & Laplace Transformation (Page 2)")
    
    p = doc.add_paragraph("Substituting the non-dimensional variables into equation (1):")
    
    eq_sub_latex = r"\frac{\nu^2}{h^3} \frac{\partial u}{\partial t} = \nu_{\mathrm{eff}} \frac{\nu}{h^3} \frac{\partial^2 u}{\partial y^2} - \frac{u \nu^2}{h \bar{\bar{K}}} + G"
    add_display_equation(eq_sub_latex)
    
    p = doc.add_paragraph("Multiplying throughout by h³ / ν²:")
    eq_simp_latex = r"\frac{\partial u}{\partial t} = \frac{\nu_{\mathrm{eff}}}{\nu} \frac{\partial^2 u}{\partial y^2} - \frac{h^2}{\bar{\bar{K}}} u + \frac{h^3}{\nu^2} G"
    add_display_equation(eq_simp_latex)
    
    p = doc.add_paragraph("Writing in non-dimensional form:")
    
    # Equation (4)
    eq4_latex = r"\frac{\partial u}{\partial t} = \gamma \frac{\partial^2 u}{\partial y^2} - \frac{a u}{Da} + P"
    add_display_equation(eq4_latex, eq_number="4")
    
    p = doc.add_paragraph("where:")
    eq_params_latex = r"a = \cos^2\theta + K^* \sin^2\theta, \quad K^* = \frac{K_1}{K_2}, \quad \gamma = \frac{\nu_{\mathrm{eff}}}{\nu}, \quad Da = \frac{K_1}{h^2}, \quad P = \frac{h^3 G}{\nu^2}"
    add_display_equation(eq_params_latex)
    
    # Parameter description table
    p = doc.add_paragraph("Summary of dimensionless groups and parameters:")
    tbl_p = doc.add_table(rows=6, cols=3)
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_p.autofit = False
    
    headers = ["Parameter / Symbol", "Definition", "Description"]
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    
    hdr_cells = tbl_p.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        p_h = hdr_cells[i].paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_h.add_run(title)
        r.font.name = "Times New Roman"
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(r'<w:shd %s w:fill="102C57"/>' % nsdecls('w'))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        
    param_data = [
        (r"\gamma", r"\nu_{\mathrm{eff}} / \nu", "Viscosity ratio (effective Brinkman to fluid kinematic viscosity)"),
        (r"K^*", r"K_1 / K_2", "Anisotropic permeability ratio"),
        (r"a", r"\cos^2\theta + K^* \sin^2\theta", "Anisotropic permeability orientation factor"),
        (r"Da", r"K_1 / h^2", "Darcy number"),
        (r"P", r"h^3 G / \nu^2", "Dimensionless constant pressure gradient / body force parameter")
    ]
    
    for row_idx, (p_sym, p_def, p_desc) in enumerate(param_data, start=1):
        row_cells = tbl_p.rows[row_idx].cells
        for j in range(3):
            row_cells[j].width = col_widths[j]
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx % 2 == 1:
                shd = parse_xml(r'<w:shd %s w:fill="F4F6F9"/>' % nsdecls('w'))
                row_cells[j]._tc.get_or_add_tcPr().append(shd)
                
        p_0 = row_cells[0].paragraphs[0]
        p_0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_math_to_paragraph(p_0, p_sym)
        
        p_1 = row_cells[1].paragraphs[0]
        p_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_math_to_paragraph(p_1, p_def)
        
        p_2 = row_cells[2].paragraphs[0]
        p_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_2.add_run(p_desc)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        
    tblBorders = parse_xml(
        r'<w:tblBorders %s>'
        r'  <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        r'  <w:left w:val="none"/>'
        r'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        r'  <w:right w:val="none"/>'
        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        r'  <w:insideV w:val="none"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tbl_p._tbl.tblPr.append(tblBorders)

    # Laplace transform application
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Apply Laplace transform on (4):")
    
    # Equation (5)
    eq5_latex = r"s \bar{u} = \gamma \frac{d^2 \bar{u}}{dy^2} - \frac{a \bar{u}}{Da} + \frac{P}{s}"
    add_display_equation(eq5_latex, eq_number="5")
    
    p = doc.add_paragraph()
    p.add_run("where ")
    add_math_to_paragraph(p, r"s > 0")
    p.add_run(" is the Laplace parameter.")
    
    p = doc.add_paragraph("Rearranging:")
    eq_rearr_latex = r"\frac{d^2 \bar{u}}{dy^2} - \left( \frac{a}{Da} + s \right) \frac{\bar{u}}{\gamma} = -\frac{P}{s \gamma}"
    add_display_equation(eq_rearr_latex)
    
    p = doc.add_paragraph("Let:")
    eq_m2_latex = r"m^2 = \left( \frac{a}{Da} + s \right) \frac{1}{\gamma}"
    add_display_equation(eq_m2_latex)
    
    p = doc.add_paragraph("Then:")
    # Equation (6)
    eq6_latex = r"\frac{d^2 \bar{u}}{dy^2} - m^2 \bar{u} = -\frac{P}{s \gamma}"
    add_display_equation(eq6_latex, eq_number="6")

    # =============================================================
    # SECTION 3: GENERAL SOLUTION & BOUNDARY CONDITIONS (PAGE 3)
    # =============================================================
    add_heading_1("3. Laplace Transform Method: General Solution & Transformed Boundary Conditions (Page 3)")
    
    p = doc.add_paragraph("Solving equation (6) gives:")
    
    # Equation (7)
    eq7_latex = r"\therefore \bar{u} = A_1 \cosh(y m) + A_2 \sinh(y m) + \frac{P}{s \gamma m^2}"
    add_display_equation(eq7_latex, eq_number="7")
    
    p = doc.add_paragraph("Also, the first derivative is:")
    
    # Equation (8)
    eq8_latex = r"\frac{d\bar{u}}{dy} = m \left[ A_1 \sinh(y m) + A_2 \cosh(y m) \right]"
    add_display_equation(eq8_latex, eq_number="8")

    p = doc.add_paragraph("From boundary conditions (2):")
    
    # Initial Condition transformed
    p_trans_ic = doc.add_paragraph()
    p_trans_ic.add_run("•  For ")
    add_math_to_paragraph(p_trans_ic, r"t \le 0")
    p_trans_ic.add_run(", ")
    add_math_to_paragraph(p_trans_ic, r"u = 0")
    p_trans_ic.add_run(" for ")
    add_math_to_paragraph(p_trans_ic, r"0 \le y \le 1")
    
    # Transformed BCs (9)
    p_trans_bc = doc.add_paragraph()
    p_trans_bc.add_run("•  For ")
    add_math_to_paragraph(p_trans_bc, r"t > 0")
    p_trans_bc.add_run(":")
    
    eq9_latex = r"\begin{cases} \bar{u} = \beta_1 \frac{d\bar{u}}{dy} & \text{at } y = 0 \\ \bar{u} = -\beta_2 \frac{d\bar{u}}{dy} & \text{at } y = 1 \end{cases}"
    add_display_equation(eq9_latex, eq_number="9")
    
    p = doc.add_paragraph("Here, equation (9) represents the corresponding dimensionless boundary conditions in the Laplace domain.")

    # Application of BC at y=0 (Page 3)
    p = doc.add_paragraph("To derive A1 and A2, we apply boundary conditions (9) to equations (7) and (8):")
    
    p_bc0 = doc.add_paragraph()
    p_bc0.add_run("At ")
    add_math_to_paragraph(p_bc0, r"y = 0")
    p_bc0.add_run(":")
    
    eq_bc0_step1 = r"A_1 + \frac{P}{s \gamma m^2} = \beta_1 m [A_2]"
    add_display_equation(eq_bc0_step1)
    
    p = doc.add_paragraph("which implies:")
    
    # Equation (10)
    eq10_latex = r"A_1 - \beta_1 m A_2 = -\frac{P}{s \gamma m^2}"
    add_display_equation(eq10_latex, eq_number="10")

    # Application of BC at y=1 (Page 3-4)
    p_bc1 = doc.add_paragraph()
    p_bc1.add_run("At ")
    add_math_to_paragraph(p_bc1, r"y = 1")
    p_bc1.add_run(":")
    
    eq_bc1_step1 = r"A_1 \cosh(m) + A_2 \sinh(m) + \frac{P}{s \gamma m^2} = -\beta_2 m \left[ A_1 \sinh(m) + A_2 \cosh(m) \right]"
    add_display_equation(eq_bc1_step1)

    # =============================================================
    # SECTION 4: ALGEBRAIC FORMULATION OF SYSTEM (PAGE 4)
    # =============================================================
    add_heading_1("4. Grouping Terms & Linear System Formulation (Page 4)")
    
    p = doc.add_paragraph("Rearranging terms from the boundary condition at y = 1:")
    
    # Equation (11)
    eq11_latex = r"A_1 \left[ \cosh(m) + \beta_2 m \sinh(m) \right] + A_2 \left[ \sinh(m) + \beta_2 m \cosh(m) \right] = -\frac{P}{s \gamma m^2}"
    add_display_equation(eq11_latex, eq_number="11")
    
    p = doc.add_paragraph("Let:")
    eq_x1_latex = r"x_1 = \cosh(m) + \beta_2 m \sinh(m)"
    add_display_equation(eq_x1_latex)
    
    eq_x2_latex = r"x_2 = \sinh(m) + \beta_2 m \cosh(m)"
    add_display_equation(eq_x2_latex)
    
    p = doc.add_paragraph("Then:")
    
    # Equation (12)
    eq12_latex = r"A_1 x_1 + A_2 x_2 = -\frac{P}{s \gamma m^2}"
    add_display_equation(eq12_latex, eq_number="12")
    
    p = doc.add_paragraph("Solve equations (10) and (12) to calculate A1 and A2.")

    # =============================================================
    # SECTION 5: EXACT ANALYTICAL SOLUTION FOR A1 AND A2 (PAGE 5)
    # =============================================================
    add_heading_1("5. Exact Analytical Solution for A1 and A2 (Page 5)")
    
    p = doc.add_paragraph("From equations (10) and (12), we have:")
    
    eq_sys_1 = r"A_1 - \beta_1 m A_2 = -\frac{P}{s \gamma m^2}"
    add_display_equation(eq_sys_1)
    
    eq_sys_2 = r"A_1 x_1 + A_2 x_2 = -\frac{P}{s \gamma m^2}"
    add_display_equation(eq_sys_2)
    
    p = doc.add_paragraph("Let:")
    eq_x3_latex = r"x_3 = \frac{P}{s \gamma m^2}"
    add_display_equation(eq_x3_latex)
    
    p = doc.add_paragraph("The system becomes:")
    eq_s1 = r"A_1 - \beta_1 m A_2 = -x_3 \quad (\times x_1)"
    add_display_equation(eq_s1)
    
    eq_s2 = r"A_1 x_1 + A_2 x_2 = -x_3 \quad (\times 1)"
    add_display_equation(eq_s2)

    # Derivation of A2
    add_heading_2("5.1 Solution for Constant A2")
    p = doc.add_paragraph("Multiplying equation (10) by x1:")
    eq_elim_a1_1 = r"A_1 x_1 - \beta_1 m x_1 A_2 = -x_1 x_3"
    add_display_equation(eq_elim_a1_1)
    
    eq_elim_a1_2 = r"A_1 x_1 + A_2 x_2 = -x_3"
    add_display_equation(eq_elim_a1_2)
    
    p = doc.add_paragraph("Subtracting the equations:")
    eq_subtr = r"A_2 (x_2 + \beta_1 m x_1) = x_1 x_3 - x_3 = x_3 (x_1 - 1)"
    add_display_equation(eq_subtr)
    
    p = doc.add_paragraph("Solving for A2 gives:")
    eq_a2_final = r"A_2 = \frac{x_3 (x_1 - 1)}{x_2 + \beta_1 m x_1}"
    add_display_equation(eq_a2_final)

    # Derivation of A1
    add_heading_2("5.2 Solution for Constant A1")
    p = doc.add_paragraph("Also:")
    eq_elim_a2_1 = r"A_1 x_2 - \beta_1 m x_2 A_2 = -x_2 x_3"
    add_display_equation(eq_elim_a2_1)
    
    eq_elim_a2_2 = r"A_1 x_1 \beta_1 m + A_2 x_2 \beta_1 m = -x_3 \beta_1 m"
    add_display_equation(eq_elim_a2_2)
    
    p = doc.add_paragraph("Adding the two equations:")
    eq_add_a1 = r"A_1 (x_2 + x_1 \beta_1 m) = -x_3 (x_2 + \beta_1 m)"
    add_display_equation(eq_add_a1)
    
    p = doc.add_paragraph("Solving for A1 gives:")
    eq_a1_final = r"A_1 = -\frac{x_3 (x_2 + \beta_1 m)}{x_2 + x_1 \beta_1 m}"
    add_display_equation(eq_a1_final)

    # =============================================================
    # SECTION 6: COMPLETE SOLUTION SUMMARY
    # =============================================================
    add_heading_1("6. Summary of the Complete Analytical Solution")
    
    p = doc.add_paragraph("The complete analytical velocity solution in the transformed Laplace domain is:")
    
    eq_sol_summary = r"\bar{u}(y, s) = A_1 \cosh(m y) + A_2 \sinh(m y) + \frac{P}{s \gamma m^2}"
    add_display_equation(eq_sol_summary)
    
    p = doc.add_paragraph("where the analytical constants and parameter definitions are:")
    
    summary_eqs = [
        (r"A_1 = -\frac{x_3 (x_2 + \beta_1 m)}{x_2 + \beta_1 m x_1}"),
        (r"A_2 = \frac{x_3 (x_1 - 1)}{x_2 + \beta_1 m x_1}"),
        (r"x_1 = \cosh(m) + \beta_2 m \sinh(m)"),
        (r"x_2 = \sinh(m) + \beta_2 m \cosh(m)"),
        (r"x_3 = \frac{P}{s \gamma m^2}"),
        (r"m = \sqrt{\frac{1}{\gamma}\left( \frac{a}{Da} + s \right)} = \sqrt{\frac{1}{\gamma}\left( \frac{\cos^2\theta + K^* \sin^2\theta}{Da} + s \right)}"),
        (r"K^* = \frac{K_1}{K_2}, \quad Da = \frac{K_1}{h^2}, \quad \gamma = \frac{\nu_{\mathrm{eff}}}{\nu}, \quad P = \frac{h^3 G}{\nu^2}")
    ]
    
    for eq_s in summary_eqs:
        add_display_equation(eq_s)

    # Save document
    output_path = r"c:\Users\widad\Documents\Opportunity Hub\YEA_PROJECT.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
